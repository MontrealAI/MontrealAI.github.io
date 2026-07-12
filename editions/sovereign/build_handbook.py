from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

base=Path('/mnt/data/goalos-partner-masterclass-sovereign-v2')
template=Path('/mnt/data/goalos-agi-alpha-what-we-built-grand-canonical-series-2026-MASTER.docx')
out=base/'docs'/'goalos-partner-masterclass-sovereign-v2-handbook.docx'
NAVY='0B1D33'; INK='14273D'; GOLD='B78324'; BLUE='285E85'; MUTED='68798E'; IVORY='FFFDF8'; PALE='F7F2E7'; LINE='D9CFBD'; LIGHT='F8F4EA'; GREEN='1F6A4B'; RED='9B2C3E'

def clear_body(doc):
    body=doc._element.body
    for child in list(body):
        if child.tag!=qn('w:sectPr'): body.remove(child)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def border(cell,color=LINE):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for e in ('top','left','right','bottom'):
        el=borders.find(qn('w:'+e))
        if el is None: el=OxmlElement('w:'+e); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),color)

def add_page_field(p):
    r=p.add_run(); f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin'); ins=OxmlElement('w:instrText'); ins.text=' PAGE '; f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'end'); r._r.extend([f1,ins,f2])

def callout(doc,title,text,color=GOLD,fill='FFF7E4'):
    t=doc.add_table(rows=2,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,color); p=c.paragraphs[0]; r=p.add_run(title.upper()); r.bold=True; r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
    c=t.cell(1,0); shade(c,fill); border(c,color); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(text); r.font.name='Georgia'; r.font.size=Pt(10.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    doc.add_paragraph()

def matrix(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,NAVY); p=c.paragraphs[0]; r=p.add_run(h); r.bold=True; r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if ri%2: shade(cells[i],'F8F9FB')
            p=cells[i].paragraphs[0]; r=p.add_run(str(v)); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(INK); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()

def body(doc,text,bold_prefix=None):
    p=doc.add_paragraph(style='Legal Body' if 'Legal Body' in doc.styles else 'Normal')
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); r.bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    return p

def bullets(doc,items):
    for x in items:
        p=doc.add_paragraph(style='Legal Body' if 'Legal Body' in doc.styles else 'Normal'); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.16); p.add_run('• '+x)

doc=Document(template); clear_body(doc)
doc.core_properties.title='GoalOS Partner MASTERCLASS — Sovereign Edition Handbook'
doc.core_properties.subject='Executive, partner, RSI, and proof-to-capability operating handbook'
doc.core_properties.author='Vincent Boucher / GoalOS'
for sec in doc.sections:
    sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72)
    hp=sec.header.paragraphs[0]; hp.text='GOALOS PARTNER MASTERCLASS · SOVEREIGN EDITION'; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: r.font.name='Arial'; r.font.size=Pt(7.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    fp=sec.footer.paragraphs[0]; fp.text='PROOF-CARRYING AUTONOMY · PARTNER HANDBOOK'; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs: r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(MUTED)
    p=sec.footer.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Page '); add_page_field(p)
    for r in p.runs: r.font.name='Arial'; r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(MUTED)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('MONTREAL.AI RESEARCH  ◇  GOALOS PARTNER SERIES'); r.font.name='Georgia'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GOLD)
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('GoalOS Partner\nMASTERCLASS'); r.font.name='Georgia'; r.font.size=Pt(34); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Sovereign Edition'); r.font.name='Georgia'; r.font.size=Pt(16); r.font.italic=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('The institution that makes intelligence earn authority'); r.font.name='Georgia'; r.font.size=Pt(14); r.font.italic=True; r.font.color.rgb=RGBColor.from_string(NAVY)
doc.add_paragraph(); doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Vincent Boucher'); r.font.name='Georgia'; r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('President, QUEBEC.AI & MONTREAL.AI'); r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(MUTED)
doc.add_paragraph(); callout(doc,'Core empirical law','Mission 1 must not merely be remembered. It must earn the right to make Mission 2 measurably better under equal constraints.',GOLD,'FFF7E4')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('AI creates output. GoalOS creates proof.'); r.font.name='Georgia'; r.font.size=Pt(11); r.font.italic=True; r.font.color.rgb=RGBColor.from_string(GOLD)
doc.add_page_break()

# Executive thesis
doc.add_heading('Executive Thesis',level=1)
body(doc,'AI capability is becoming abundant. Institutional permission to rely on it is not. GoalOS is the proof-to-capability operating layer that determines which autonomous work may become trusted, paid, remembered, reused, challenged, revoked, or permitted to improve future work.')
callout(doc,'Institutional distinction','A prompt asks a model to answer. A proof loop decides what must be proven before the answer can matter.',BLUE,'EDF4FA')
body(doc,'The canonical loop is:')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Objective → Mission Contract → Proof Debt → AGI Jobs → ProofBundles → Evidence Docket → AGI Node Validation → Chronicle → Validated Skill → Merkle Graph Root → Mission 2'); r.font.name='Arial'; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
body(doc,'The flagship deliverable is a Governed Decision State: not another static report, but an inspectable, challengeable, rollback-ready state containing evidence, validation, risk, action, memory, and reusable capability.')
matrix(doc,['Partner','Primary value','Recommended entry'],[
('Executive / Board','Decision readiness and defensibility','Proof Mission Sprint'),('Chief AI / Technology','Governed capability transfer','Enterprise Proof OS'),('Research / Frontier Lab','Reproduced research advantage','RSI Evidence Docket'),('Risk / Assurance / Legal','Evidence, responsibility, rollback','Validation Council'),('Infrastructure / Cloud / Data','Trusted execution and replay capacity','AGI Node Council'),('Strategic Capital / Ecosystem','Verified-work throughput and reusable capability','Cross-Institution Proof Network')])

# Module map
doc.add_heading('The 18-Module Institution',level=1)
modules=[
('00','Partner Orientation','Choose role, pace, organization, and consequential objective.'),('01','The Category','Place GoalOS above models and agents as the authority layer.'),('02','Mission OS Studio','Freeze objective, proof level, risk, privacy, validators, blocked claims, rollback.'),('03','Proof Factory','Convert unsupported claims into Proof Debt and custom AGI Jobs.'),('04','Evidence Docket','Make work, failures, sources, replay, risks, and blocked claims inspectable.'),('05','AGI Node Validation Council','Use independent commit-reveal validation and correlation detection.'),('06','Chronicle & Validated Skill Graph','Admit only scoped, evidenced, replayable, versioned, rollbackable capability.'),('07','Merkle, Privacy & ZK Proof Room','Commit exact state without plaintext private intelligence; revoke on tamper.'),('08','Bonded Authority Economy','Secure payment, trust, challenge, reuse, nodes, and upgrades with capital at risk.'),('09','RSI Sovereign Invention Lab','Run recursive improvement through baseline, evidence, stress, persistence, and rollback.'),('10','Move-37 Breakthrough Control','Apply higher skepticism to high-novelty candidates.'),('11','The Decisive Mission 1 → Mission 2 Test','Compare fresh, raw-memory, validated-skill, and ungated-memory arms.'),('12','REAL-001 Evidence Room','Inspect a consequential public GitHub maintenance transfer case.'),('13','Business Value & Proof Economics','Model Proof Debt exposure, cycle-time release, loss reduction, and reuse.'),('14','Partner Architecture Mapper','Map GoalOS into enterprise, research, infrastructure, assurance, or capital stacks.'),('15','30-Day Pilot Architect','Generate a partner-specific charter and Mission 2 metric.'),('16','Trust, Due Diligence & Claim Boundary','Separate evidence, simulations, pending gates, and prohibited overclaims.'),('17','Executive Close & Partnership Dossier','Generate the brief, completion record, certificate, and next step.')]
matrix(doc,['No.','Module','Outcome'],modules)
doc.add_page_break()

# Modules details grouped.
details={
'00 Partner Orientation':('Personalize the experience',['Choose the authority problem, not merely the technology category.','Select the proof burden and partner perspective.','Define the consequential objective and executive sponsor.'],'Output: role-adaptive thesis and recommended pathway.'),
'01 The Category':('Separate capability from authority',['Models produce candidate output.','Agents perform bounded action.','GoalOS grants earned authority through proof, validation, memory, settlement, challenge, and rollback.'],'Output: category map and partner positioning.'),
'02 Mission OS Studio':('Freeze the institutional boundary',['Objective and decision to support.','Proof level, risk class, privacy mode, and validators.','Success, failure, blocked claims, done condition, and rollback.'],'Output: signed or hashed Mission Contract.'),
'03 Proof Factory':('Turn uncertainty into work',['Extract high-impact unsupported claims.','Create only mission-specific proof contracts.','Give every job a worker, validator, sentinel, deliverables, acceptance tests, blocked claims, and return path.'],'Output: Proof Debt Register and AGI Job Manifest.'),
'04 Evidence Docket':('Create the proof room',['Tie claims to primary evidence, ProofBundles, contradictions, cost, risk, and replay.','Preserve failures and negative evidence.','Make blocked claims and uncertainty visible.'],'Output: public-safe Evidence Docket.'),
'05 AGI Node Validation Council':('Make validation effectively independent',['Direct validator identities and separate signing keys.','Commit-reveal to reduce copying and strategic voting.','Diversity across operator, model, cloud, region, jurisdiction, and wallet control.'],'Output: signed validator record and independence assessment.'),
'06 Chronicle & VSG':('Convert evidence into bounded memory',['Chronicle admits, narrows, repairs, rejects, quarantines, supersedes, revokes, or retires.','Validated Skills preserve scope, method, evidence, tests, validators, risk, freshness, utility, lineage, and rollback.'],'Output: Chronicle decision and Skill Passport.'),
'07 Merkle, Privacy & ZK':('Commit exact state without exposing private intelligence',['Typed roots for skills, evidence, bundles, policies, attestations, and lineage.','Encrypted or off-chain private methods.','Root mismatch immediately revokes inherited authority.'],'Output: graph epoch root and tamper demonstration.'),
'08 Bonded Authority':('Make consequence economically accountable',['Worker, validator, claim, challenger, reuse, node, proposal, canary, and rollback bonds.','No bond, no authority transition.','Customer token interaction can remain hidden behind professional protocol security.'],'Output: illustrative bond ledger and settlement gate.'),
'09 RSI Sovereign Invention Lab':('Govern recursive improvement',['TARGET → EMIT → FILTER → ATLAS → TEST-PLAN → EVAL → INSERT → PROMOTE.','Temperature-zero manifests, schema validation, baselines, append-only ledgers.','Search allocates exploration; mechanical gates decide promotion.'],'Output: replayable RSI state and decision.'),
'10 Move-37 Control':('Treat novelty as a reason for more skepticism',['Recognize thresholds and evidence contact.','Reproduce with fixed seeds and baseline.','Stress under policy shocks.','Require persistence and package a dossier.'],'Output: breakthrough-admission dossier or hold.'),
'11 Decisive Transfer Test':('Measure compounding rather than memory',['Fresh control, raw memory, Validated Skill, and ungated candidate.','Same model, tools, time, compute, evaluator, stopping rule, and prior budget.','Fresh held-out work.'],'Output: Mission 2 capability-transfer report.'),
'12 REAL-001':('Inspect a consequential real-task mechanism',['Validated Skill score 98.57; fresh and raw 92.86; ungated candidate -15.00.','Correct safe decision: controlled migration required.','First-party reconstruction passed; independent external replay remains pending.'],'Output: Proof Pack, Evidence Docket, replay receipt, and workflows.'),
'13 Business Value':('Translate proof into enterprise value',['Proof Debt exposure.','Expected loss reduction.','Decision-cycle days released.','Missions receiving reusable capability.'],'Output: claim-bounded scenario model.'),
'14 Partner Architecture':('Map into existing systems',['Existing models and agents remain workers.','GoalOS freezes authority, evidence, validation, memory, and commitment.','Partner pathways differ for enterprise, research, infrastructure, assurance, validators, and capital.'],'Output: integration map and recommended entry.'),
'15 30-Day Pilot':('Move from briefing to evidence',['Week 1 freeze authority.','Week 2 produce proof.','Week 3 validate and challenge.','Week 4 deliver and test transfer.'],'Output: partner-specific pilot charter.'),
'16 Trust & Due Diligence':('Make evidence maturity part of the product',['State what is demonstrated, simulated, pending, and prohibited.','Keep independent replay, delayed outcomes, production authority, token value, legal certification, and general RSI separately gated.'],'Output: due-diligence matrix and claim boundary.'),
'17 Executive Close':('Convert interest into a bounded next step',['One consequential objective.','One executive sponsor.','One independent reviewer.','One measured Mission 2 test.'],'Output: executive dossier and completion record.')}
for idx,(name,(thesis,items,outcome)) in enumerate(details.items()):
    doc.add_heading(name,level=1)
    callout(doc,'Module thesis',thesis,BLUE,'EDF4FA')
    bullets(doc,items)
    body(doc,outcome,bold_prefix='Output:')
    if idx%2==1 and idx<len(details)-1: doc.add_page_break()

# Formats and pilot
doc.add_page_break(); doc.add_heading('Presentation Formats',level=1)
matrix(doc,['Format','Use','Core modules'],[
('Executive 15 minutes','Category and partnership decision','01, 02, 05, 06, 11, 12, 15, 17'),('Boardroom 30 minutes','Strategic diligence and value','Add 04, 07, 08, 13, 16'),('Working session 90 minutes','Build a real partner mission','Complete all operating modules'),('Self-paced certification','Deep institutional understanding','All 18 modules + dossier')])
callout(doc,'Opening line','AI capability is becoming abundant. Institutional permission to rely on it is not. GoalOS makes autonomous work earn authority.',GOLD,'FFF7E4')
callout(doc,'Closing line','Bring one consequential objective, one executive sponsor, and one independent reviewer. In 30 days, produce the first Governed Decision State and test whether the first admitted capability makes Mission 2 measurably better.',GOLD,'FFF7E4')

doc.add_heading('Thirty-Day Partner Pilot',level=1)
matrix(doc,['Week','Authority state','Deliverables'],[
('1','Freeze authority','Mission Contract, proof level, claim boundary, baseline, privacy, validators, rollback, Mission 2 metric'),('2','Produce proof','AGI Jobs, ProofBundles, provenance, contradictions, cost/risk ledgers, Evidence Docket'),('3','Validate and challenge','Independent replay, commit-reveal, falsification, Chronicle recommendation'),('4','Deliver and transfer','Governed Decision State, Skill Passport, Merkle packet, Mission 2 comparison, executive dossier')])

doc.add_heading('Claim Boundary',level=1)
body(doc,'The masterclass combines interactive simulations, reference architecture, and REAL-001 first-party real-task evidence. It does not establish achieved AGI or ASI, general real-world recursive self-improvement, empirical state of the art, independent external validation, production certification, guaranteed outcomes, legal or tax advice, token value, liquidity, yield, regulatory exemption, or autonomous production authority.')
callout(doc,'Final operating law','Generate freely. Prove rigorously. Validate independently. Admit cautiously. Commit exactly. Reuse within scope. Challenge economically. Roll back decisively. Improve only through evidence.',NAVY,'EDF4FA')

doc.save(out)
print(out)
