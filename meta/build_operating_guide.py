from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT=Path('/mnt/data/GoalOS_Partner_MASTERCLASS_MASTERPIECE_v3')
A=OUT/'assets'
DOC=OUT/'guides/GoalOS_Partner_MASTERCLASS_MASTERPIECE_Operating_Guide.docx'
NAVY='0B1E34'; GOLD='B9872C'; GOLD2='E7C978'; BLUE='337DAE'; CYAN='2A95A6'; GREEN='2D8B63'; RED='A63B51'; PURPLE='66529C'; INK='182B43'; MUTED='66778C'; IVORY='F7F3EA'; LIGHT='EEF3F7'; WHITE='FFFFFF'; LINE='D7DEE6'

doc=Document();sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=Inches(.55);sec.bottom_margin=Inches(.55);sec.left_margin=Inches(.62);sec.right_margin=Inches(.62);sec.header_distance=Inches(.22);sec.footer_distance=Inches(.25)
styles=doc.styles
styles['Normal'].font.name='Arial';styles['Normal'].font.size=Pt(9.2);styles['Normal'].font.color.rgb=RGBColor.from_string(INK);styles['Normal'].paragraph_format.space_after=Pt(4);styles['Normal'].paragraph_format.line_spacing=1.05
for nm,size,color,space in [('Title',31,NAVY,8),('Subtitle',14,MUTED,7),('Heading 1',24,NAVY,9),('Heading 2',14,GOLD,5),('Heading 3',11,NAVY,3)]:
 st=styles[nm];st.font.name='Georgia' if nm in ('Title','Heading 1','Heading 2') else 'Arial';st.font.size=Pt(size);st.font.bold=nm!='Subtitle';st.font.color.rgb=RGBColor.from_string(color);st.paragraph_format.space_before=Pt(0);st.paragraph_format.space_after=Pt(space);st.paragraph_format.keep_with_next=True
if 'Kicker' not in styles:
 st=styles.add_style('Kicker',WD_STYLE_TYPE.PARAGRAPH);st.font.name='Arial';st.font.size=Pt(7.5);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(GOLD);st.paragraph_format.space_after=Pt(5);st.paragraph_format.keep_with_next=True
if 'Pull Quote' not in styles:
 st=styles.add_style('Pull Quote',WD_STYLE_TYPE.PARAGRAPH);st.font.name='Georgia';st.font.size=Pt(13);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(NAVY);st.paragraph_format.left_indent=Inches(.18);st.paragraph_format.right_indent=Inches(.18);st.paragraph_format.space_before=Pt(6);st.paragraph_format.space_after=Pt(7)
if 'Small' not in styles:
 st=styles.add_style('Small',WD_STYLE_TYPE.PARAGRAPH);st.font.name='Arial';st.font.size=Pt(7.4);st.font.color.rgb=RGBColor.from_string(MUTED);st.paragraph_format.space_after=Pt(3)

def shade(cell,color):
 tcPr=cell._tc.get_or_add_tcPr();shd=tcPr.find(qn('w:shd'))
 if shd is None:shd=OxmlElement('w:shd');tcPr.append(shd)
 shd.set(qn('w:fill'),color)
def margins(cell,top=80,start=100,bottom=80,end=100):
 tc=cell._tc;tcPr=tc.get_or_add_tcPr();tcMar=tcPr.first_child_found_in('w:tcMar')
 if tcMar is None:tcMar=OxmlElement('w:tcMar');tcPr.append(tcMar)
 for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
  node=tcMar.find(qn('w:'+m))
  if node is None:node=OxmlElement('w:'+m);tcMar.append(node)
  node.set(qn('w:w'),str(v));node.set(qn('w:type'),'dxa')
def set_cell_text(cell,text,bold=False,color=INK,size=8.5,align=None):
 cell.text='';p=cell.paragraphs[0];p.paragraph_format.space_after=Pt(0);r=p.add_run(text);r.bold=bold;r.font.name='Arial';r.font.size=Pt(size);r.font.color.rgb=RGBColor.from_string(color)
 if align is not None:p.alignment=align
 cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;margins(cell)
def no_borders(table):
 tblPr=table._tbl.tblPr;b=tblPr.first_child_found_in('w:tblBorders')
 if b is None:b=OxmlElement('w:tblBorders');tblPr.append(b)
 for edge in ('top','left','bottom','right','insideH','insideV'):
  el=OxmlElement('w:'+edge);el.set(qn('w:val'),'nil');b.append(el)
def border_table(table,color=LINE,size='4'):
 tblPr=table._tbl.tblPr;b=tblPr.first_child_found_in('w:tblBorders')
 if b is None:b=OxmlElement('w:tblBorders');tblPr.append(b)
 for edge in ('top','left','bottom','right','insideH','insideV'):
  el=OxmlElement('w:'+edge);el.set(qn('w:val'),'single');el.set(qn('w:sz'),size);el.set(qn('w:color'),color);b.append(el)
def field(paragraph,code):
 r=paragraph.add_run();fldChar=OxmlElement('w:fldChar');fldChar.set(qn('w:fldCharType'),'begin');instr=OxmlElement('w:instrText');instr.set(qn('xml:space'),'preserve');instr.text=code;sep=OxmlElement('w:fldChar');sep.set(qn('w:fldCharType'),'separate');txt=OxmlElement('w:t');txt.text='1';end=OxmlElement('w:fldChar');end.set(qn('w:fldCharType'),'end');r._r.extend([fldChar,instr,sep,txt,end])

def header_footer():
 h=sec.header;p=h.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.text='GOALOS PARTNER MASTERCLASS  ·  MASTERPIECE INSTITUTIONAL EDITION';p.style='Kicker';p.runs[0].font.size=Pt(6.5)
 f=sec.footer;t=f.add_table(rows=1,cols=2,width=Inches(7.25));t.columns[0].width=Inches(5.7);t.columns[1].width=Inches(1.55);no_borders(t);set_cell_text(t.cell(0,0),'AI creates output. GoalOS creates proof.',False,MUTED,7.2);p=t.cell(0,1).paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;r=p.add_run('PAGE ');r.font.name='Arial';r.font.size=Pt(7.2);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(MUTED);field(p,'PAGE')
header_footer()

def add_kicker(text):doc.add_paragraph(text.upper(),'Kicker')
def add_title(text,sub=None):doc.add_heading(text,0 if len(doc.paragraphs)<2 else 1); 

def heading(text,kicker=None,sub=None):
 if kicker:add_kicker(kicker)
 doc.add_heading(text,1)
 if sub:
  p=doc.add_paragraph(sub,'Subtitle');p.paragraph_format.space_after=Pt(7)
def para(text,bold_lead=None,style=None):
 p=doc.add_paragraph(style=style)
 if bold_lead and text.startswith(bold_lead):
  r=p.add_run(bold_lead);r.bold=True;p.add_run(text[len(bold_lead):])
 else:p.add_run(text)
 return p
def bullets(items,level=0):
 for item in items:
  p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2');p.paragraph_format.space_after=Pt(2);p.paragraph_format.left_indent=Inches(.2+.18*level);p.paragraph_format.first_line_indent=Inches(-.12);p.add_run(item)
def quote(text,color=GOLD):
 t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;t.columns[0].width=Inches(7.05);shade(t.cell(0,0),IVORY);margins(t.cell(0,0),140,180,140,180);border_table(t,color,'8');p=t.cell(0,0).paragraphs[0];p.style='Pull Quote';p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run(text)
 doc.add_paragraph().paragraph_format.space_after=Pt(0)
def image(name,width=7.0,caption=None):
 p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run();r.add_picture(str(A/name),width=Inches(width));p.paragraph_format.space_after=Pt(3)
 if caption:
  c=doc.add_paragraph(caption,'Small');c.alignment=WD_ALIGN_PARAGRAPH.CENTER;c.runs[0].italic=True

def two_col(items,headers=None,widths=(2.2,4.8),header_color=NAVY):
 t=doc.add_table(rows=1 if headers else 0,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;t.columns[0].width=Inches(widths[0]);t.columns[1].width=Inches(widths[1])
 if headers:
  for j,h in enumerate(headers):shade(t.cell(0,j),header_color);set_cell_text(t.cell(0,j),h,True,WHITE,8.2)
 for i,(a,b) in enumerate(items):
  cells=t.add_row().cells;shade(cells[0],LIGHT if i%2==0 else WHITE);shade(cells[1],LIGHT if i%2==0 else WHITE);set_cell_text(cells[0],a,True,NAVY,8.5);set_cell_text(cells[1],b,False,INK,8.4)
 border_table(t);doc.add_paragraph().paragraph_format.space_after=Pt(0);return t

def cards(items,cols=3):
 rows=(len(items)+cols-1)//cols;t=doc.add_table(rows=rows,cols=cols);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
 for c in range(cols):t.columns[c].width=Inches(7.1/cols)
 for i in range(rows*cols):
  cell=t.cell(i//cols,i%cols);margins(cell,120,130,120,130)
  if i>=len(items):shade(cell,WHITE);continue
  title_,body_,color=items[i];shade(cell,LIGHT);p=cell.paragraphs[0];p.paragraph_format.space_after=Pt(4);r=p.add_run(title_);r.bold=True;r.font.name='Georgia';r.font.size=Pt(10.5);r.font.color.rgb=RGBColor.from_string(color);p2=cell.add_paragraph();p2.paragraph_format.space_after=Pt(0);r=p2.add_run(body_);r.font.name='Arial';r.font.size=Pt(7.8);r.font.color.rgb=RGBColor.from_string(INK)
 border_table(t);doc.add_paragraph().paragraph_format.space_after=Pt(0)
def page():doc.add_page_break()

# Cover
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(12);r=p.add_run('GOALOS');r.font.name='Arial';r.font.size=Pt(11);r.bold=True;r.font.color.rgb=RGBColor.from_string(GOLD)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Partner MASTERCLASS\nMASTERPIECE Edition');r.font.name='Georgia';r.font.size=Pt(31);r.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('The complete partner institution for proof-carrying intelligence');r.italic=True;r.font.name='Georgia';r.font.size=Pt(14);r.font.color.rgb=RGBColor.from_string(MUTED)
image('apex_atrium.png',6.35)
quote('A model may answer. An agent may act. An institution must prove. Open-ended work may be attempted; strictly gated in what authority a result may earn.')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Release GOS-PM-MASTERPIECE-2026.07-v3.0  ·  July 2026');r.font.size=Pt(7.5);r.font.color.rgb=RGBColor.from_string(MUTED)
page()

heading('What the Masterpiece Edition is','Executive synthesis','A presentation system, operating laboratory, evidence track, diligence room, facilitator system and first-partner conversion architecture in one coordinated release.')
cards([('4 coordinated editions','APEX, Grand, Sovereign Evidence and Production Showcase remain intact and independently usable.',GOLD),('18 deep modules','The curriculum moves from proof gap through Mission 2 and partner charter.',BLUE),('10+ operating labs','Mission, Docket, Council, stress, Merkle, RSI, diligence and deal architecture.',PURPLE),('1 decisive test','Mission 1 must earn the right to make a fresh Mission 2 measurably better.',GREEN)],2)
para('The ceiling moved. The strongest masterclass is not the longest monolith. It is a layered institution that moves the same audience from comprehension to operation, from operation to evidence, and from evidence to a bounded partnership decision.','The ceiling moved.')
quote('The partner should leave with one named objective, one decision owner, one reviewer, one evidence package, one rollback condition and one fresh Mission 2 test.')
page()

heading('The four front doors','01 · Masterclass architecture')
image('launcher.png' if (A/'launcher.png').exists() else 'grand_atrium.png',6.9,'Use the Masterpiece launcher as the orchestration layer; enter a source edition only when the room needs its depth.')
two_col([('APEX Boardroom','Premium 90-second, 12-minute and 45-minute executive theatre.'),('Grand Institution','Complete 90-minute, three-hour and one-day masterclass with labs, AI Council, diligence and Deal Architect.'),('Sovereign Evidence','Eighteen modules, REAL-001, replay, governed RSI and Mission 1 → Mission 2 transfer.'),('Production Showcase v2.1','Concise commercial mission, Evidence Docket, sample Governed Decision State and calibrated receipt.')],('Edition','Best use'))
page()

heading('Choose the delivery mode before the content','02 · Session routing')
two_col([('90 seconds','Create the category and ask which consequential decision deserves proof.'),('12 minutes','Run the live partner mission and nominate one Proof Mission.'),('45 minutes','Combine executive theatre, one live gate and REAL-001.'),('90 minutes','Draft a Mission Contract and Partner Charter.'),('Three hours','Operate the full proof loop and design an executable engagement.'),('One day','Produce a 30/60/90 institution blueprint and diligence plan.')],('Mode','Output'))
quote('The success metric is not “slides completed.” It is “a real decision acquired an explicit proof path.”')
page()

heading('The 90-second boardroom opening','03 · Executive theatre')
para('0:00–0:15 — Name the scarcity. AI output is abundant. Institutional authority is scarce.')
para('0:15–0:40 — Define the category. GoalOS is the Proof OS between autonomous work and the authority to trust, pay, remember, reuse or propagate it.')
para('0:40–1:05 — Show the deliverable. The endpoint is a Governed Decision State: claims, evidence, contradictions, risk, action, memory and rollback.')
para('1:05–1:30 — Ask for one mission. “Which consequential but reversible decision should we freeze into a Mission Contract?”')
quote('Do not begin with token mechanics, planetary scale or a long-range AGI thesis. Begin with the proof gap the room already recognizes.')
page()

heading('The 12-minute partner demonstration','04 · Commercial path')
image('showcase_overview.png',6.8,'Production Showcase v2.1 is the shortest route from category to a candidate Proof Mission.')
two_col([('1:30','Frame the proof gap and product category.'),('5:00','Run Strategic Partnership Diligence through the nine-stage mission state machine.'),('3:00','Inspect passed, conditional, disputed and blocked claims in the Evidence Docket.'),('2:30','Name sponsor, reviewer, evidence package, rollback and decision date.')],('Time','Action'))
page()

heading('The 45-minute high-caliber session','05 · Board understanding + evidence')
two_col([('5 minutes','APEX category theatre: why proof is the missing enterprise layer.'),('10 minutes','Mission Contract and Proof Debt: freeze the decision before work begins.'),('10 minutes','REAL-001: show the scoped transfer case and open Proof Debt.'),('10 minutes','Governed RSI: evidence contact, baseline, replay, persistence and independent review.'),('10 minutes','Partner Deal Architect: convert interest into a bounded pilot memorandum.')],('Segment','Purpose'))
quote('The 45-minute room should end with sponsor interest and one candidate decision—not a request to “learn more.”')
page()

heading('The 90-minute operating masterclass','06 · Mission to partner charter')
cards([('Executive theorem','Category, authority stack and commercial relevance.',GOLD),('Mission + Proof Factory','Mission Contract, Proof Debt, custom AGI Jobs and ProofBundles.',BLUE),('Docket + Council','Claims, contradictions, reviewer independence and commit-reveal.',CYAN),('Chronicle + VSG','What may enter memory, under what scope and rollback.',GREEN),('Mission 2 transfer','Use REAL-001 to show what a compounding claim requires.',PURPLE),('Partner Charter','Draft sponsor, reviewer, decision gate and expansion conditions.',RED)],3)
para('Facilitator discipline: stop after each gate and ask, “What authority did that evidence earn—and what remains blocked?”')
page()

heading('The three-hour institution-design workshop','07 · Complete operating loop')
two_col([('20 minutes','Proof OS category, product definition and authority stack.'),('35 minutes','Mission Contract workshop on the partner’s real decision.'),('35 minutes','Proof Factory and Evidence Docket lab.'),('25 minutes','AI Council, contradictions, risk and stress theatre.'),('25 minutes','Chronicle, Validated Skill Graph, Merkle and private proof.'),('20 minutes','Governed RSI, REAL-001 and Mission 2 protocol.'),('20 minutes','Partner Deal Architect and 30/60/90 charter.')],('Duration','Work product'))
quote('Every module must produce an artifact. Passive comprehension is not mastery.')
page()

heading('The one-day Partner Institution Lab','08 · Full engagement design')
para('Morning — Establish the category, select a real objective, freeze the Mission Contract and map all Proof Debt.')
para('Midday — Create mission-specific AGI Jobs, ProofBundle requirements, the claims matrix, validator route and public/private evidence boundary.')
para('Afternoon — Run Chronicle, VSG, Merkle, bonded-authority and governed-RSI labs; design the fresh Mission 2 transfer test.')
para('Closing — Complete the diligence checklist, participation record, Founding Proof Mission memorandum and 90-day institution build roadmap.')
cards([('Output 1','Signed mission thesis and named decision owner.',GOLD),('Output 2','Reviewer-ready proof and challenge architecture.',CYAN),('Output 3','Partner Charter with commercial and evidence gates.',GREEN),('Output 4','Mission 2 protocol that can falsify compounding.',PURPLE)],2)
page()

heading('The category: Proof OS for autonomous work','09 · Strategic thesis')
two_col([('Report-only AI','Output is the endpoint; evidence is reconstructed after the fact.'),('GoalOS','A Governed Decision State is the endpoint; evidence, risk, action and rollback travel together.'),('Report-only memory','Persuasive output may become an institutional prior.'),('Chronicle-gated memory','Only admitted, scoped and replayable capability may influence future missions.'),('Report-only “improvement”','A better-looking answer is treated as progress.'),('Governed improvement','Mission 2 must show equal-constraint, replayable, risk-bounded advantage.')],('Old pattern','GoalOS pattern'))
quote('A prompt asks a model to answer. A proof loop decides what must be established before the answer can matter.')
page()

heading('The canonical proof-to-capability loop','10 · Operating law')
image('proof_loop.png',4.45,'Authority appears only after the next gate passes.')
para('Objective → Mission Contract → Proof Debt → Custom AGI Jobs → ProofBundles → Evidence Docket → independent validation → Chronicle → Validated Skill → Merkle epoch → bonded authority → harder Mission 2 → delayed outcome → governed improvement.')
quote('No proof, no evolution. No evaluation, no propagation. No rollback, no release.')
page()

heading('Mission Contract: freeze authority before work','11 · The first control object')
two_col([('Decision','What exact choice must the institution make?'),('Success / failure','What evidence would justify proceed, hold, reject or repair?'),('Scope and boundary','What may the mission access or change—and what remains prohibited?'),('Proof level','Internal, external-review, controlled handoff or Chronicle-ready?'),('Review authority','Who may approve, challenge, block or require repair?'),('Done condition','Which artifacts and gates must exist before the run can end?'),('Rollback','What prior state is restored if assumptions fail?')],('Field','Institutional question'))
quote('No ambiguous authority. A mission cannot acquire broader permission merely because the work appears impressive.')
page()

heading('Proof Debt and custom AGI Jobs','12 · Uncertainty becomes work')
para('Proof Debt is the set of unsupported claims, missing artifacts, unresolved contradictions and untested assumptions that prevent a decision from earning authority.')
two_col([('Source claim','The exact claim whose authority is requested.'),('Why needed','Why the decision cannot close without this proof.'),('Worker / tools','Who or what may produce the evidence, under which boundary.'),('Validator','Who must independently review the result.'),('Deliverables','Artifacts, logs, manifests, tests, cost and risk ledgers.'),('Acceptance tests','Mechanical and expert criteria for pass, repair or block.'),('Failure branch','What happens when evidence is absent or contradictory.'),('Rollback','How any candidate action is reversed.')],('AGI Job field','Requirement'))
quote('The job factory begins empty. Jobs appear only after the frozen mission exposes actual Proof Debt.')
page()

heading('ProofBundles and the Evidence Docket','13 · Reviewer-ready proof')
cards([('ProofBundle','Artifacts, environment pins, inputs, outputs, tests, signatures, cost, risk and replay.',BLUE),('Claims matrix','Every major claim is passed, conditional, disputed, blocked or unresolved.',GOLD),('Contradiction register','Conflicting facts and assumptions remain visible and decision-relevant.',RED),('Verifier report','Reviewer identity, independence, verdict, challenge and limitation.',GREEN),('Claim boundary','Exactly what the Docket supports—and what it does not.',PURPLE),('Action Graph','Owners, dependencies, approvals, proof requirements, stops and rollback.',CYAN)],3)
quote('A proof page is not a marketing page. A qualified reviewer must be able to reconstruct what happened and what remains blocked.')
page()

heading('AGI Node Council: independence, not headcount','14 · Validation')
image('ui_council.png',4.7,'Committee routing should measure effective independence across control and epistemic dimensions.')
two_col([('Operator / wallet','Are reviewers controlled by distinct persons and economic interests?'),('Model family','Is the committee epistemically diverse?'),('Cloud / hardware','Could one infrastructure failure corrupt the entire council?'),('Jurisdiction / region','Are legal and operational dependencies correlated?'),('Capability','Does the council cover the actual proof class?'),('Commit-reveal','Were verdicts committed before other reviewers revealed?')],('Dimension','Question'))
page()

heading('Chronicle: the memory firewall','15 · Governed memory')
cards([('Admit','The candidate may influence future work within exact proven scope.',GREEN),('Admit with scope','Use is limited to named conditions, environments or risk classes.',CYAN),('Repair','Additional proof is required before admission.',GOLD),('Reject','Evidence is insufficient or advantage is absent.',RED),('Quarantine','Risk, contradiction or provenance remains unresolved.',PURPLE),('Supersede / revoke / retire','Authority can shrink or disappear when evidence changes.',MUTED)],3)
quote('No Chronicle entry, no future-mission influence.')
page()

heading('Validated Skill Graph: the compounding asset','16 · Capability memory')
image('skill_graph.png',4.45,'A validated capability passport carries scope, evidence, replay, freshness and rollback.')
bullets(['Stable identity and immutable version.','Allowed and excluded uses.','Method, evidence, tests and validator history.','Replay path, risk posture and freshness.','Rollback target, revocation and supersession.','Mission lineage and measured future-mission contribution.'])
quote('The moat is not prompt text. It is accumulated method, evidence, scope, validator history, delayed outcomes, policy and lineage.')
page()

heading('Merkle commitments and private intelligence','17 · Public-private proof boundary')
image('aep_boundary.png',4.5,'Publish the minimum proof required for trust; keep sensitive execution and context controlled.')
two_col([('Public proof','Hashes, Merkle roots, attestations, selection certificates, receipts, revocations and public-safe Docket manifests.'),('Private intelligence','Prompts, long traces, customer data, proprietary tools, evaluation workpapers and privileged rationale.'),('Chronicle','Decides whether capability earns authority.'),('Merkle root','Commits the exact admitted graph state; inclusion proves membership and integrity, not truth.')],('Plane','Contents and authority'))
page()

heading('$AGIALPHA: bonded authority, not the partner opening','18 · Economic accountability')
two_col([('Work begins','Mission capacity and Worker Proof Bond.'),('Work becomes paid','Accepted ProofBundle and Validator Accuracy Bonds.'),('Claim becomes trusted','Claim Authority Bond, Chronicle and challenge finality.'),('Claim is challenged','Challenger Bond and evidence-bearing resolution.'),('Skill becomes reusable','Capability Reuse Bond, delayed outcome, scope and freshness.'),('Node validates','Operational reserve, reputation and independence.'),('GoalOS upgrades','Proposal, evaluator, canary and rollback reserves.')],('Authority transition','Required security'))
para('Presentation rule: lead with proof and enterprise value. Introduce $AGIALPHA only after the room understands why economically consequential authority transitions require bonded accountability.')
page()

heading('Governed recursive self-improvement','19 · RSI control plane')
image('rsi_kernel.png',6.55,'TARGET → EMIT → FILTER → ATLAS → TEST-PLAN → EVAL → INSERT → PROMOTE')
bullets(['Exploration is allowed; outcome authority is mechanical.','Interestingness may allocate search but cannot approve promotion.','Promotion requires executed evidence, baseline advantage, replay, risk control, persistence and validation.','State must be append-oriented and drift-detectable.','Every promoted candidate must carry scope, monitor and rollback.'])
page()

heading('Evidence Contact Index','20 · Confidence cannot inflate without execution')
two_col([('E0 — Simulated','LLM-only reasoning or scenario generation; idea status only.'),('E1 — Probed','Schema, static analysis or cheap microbench; probe status.'),('E2 — Executed','Run in a sandbox, benchmark, tool or environment; required beyond probe.'),('E3 — Replayed','Independent reproduction from pinned seeds, dependencies and manifests.'),('E4 — Stress-tested','Advantage survives shocks, perturbations, alternate seeds and nearby baselines.'),('E5 — Externally validated','External expert, real outcome, market settlement or delayed evidence confirms the result.')],('Level','Promotion implication'))
quote('High confidence without evidence contact is not maturity; it is evidence inflation.')
page()

heading('Move-37 breakthrough handling','21 · High novelty raises the burden')
image('move37.png',6.3,'A breakthrough is admitted only as a reproducible, stress-tested state transition.')
two_col([('Recognize','Record novelty distance, Advantage Delta, risk, ECI, lineage and archive cell.'),('Reproduce','Rerun candidate and null, incumbent and neighbor baselines under fixed conditions.'),('Stress-test','Apply policy shocks, adversarial probes, alternate seeds and side-effect scans.'),('Persistence gate','Require positive cost- and risk-adjusted advantage after shocks and replays.'),('Dossier','Package evidence, validation, governance and replay instructions before strategic promotion.')],('Gate','Requirement'))
page()

heading('Mission 2: the decisive compounding test','22 · Transfer and persistence')
image('sovereign_transfer.png',6.65,'REAL-001 demonstrates the structure of a scoped Mission 1 → Mission 2 transfer claim.')
para('Compounding claim = positive, replayable, risk-bounded, persistent Advantage Delta on a fresh task under equal constraints.')
bullets(['Use an admitted, scoped capability from Mission 1.','Choose a fresh adjacent task that was not used to create the capability.','Run null, fresh/raw, incumbent and neighboring baselines.','Record cost, risk, steps, reviewer notes and failure modes.','Promote, repair or retire based on the measured delta.'])
quote('No Mission 2, no compounding claim.')
page()

heading('REAL-001: what it proves and what remains open','23 · First-party evidence')
cards([('+5.71','Reported score lift for the skill arm versus fresh/raw in the scoped case.',GREEN),('−1 step','One fewer observation step in the repository-specific transfer task.',CYAN),('MATCH','A first-party deterministic replay receipt reproduces the recorded result.',GOLD),('BLOCKED','The ungated arm did not receive authority.',RED)],2)
para('Supported claim: a repository-specific transfer mechanism can preserve a scoped capability and improve one fresh adjacent task under the reported conditions.')
para('Open Proof Debt: two independent operator replays and a human maintainer review of the proposed migration action. The case does not establish arbitrary-domain compounding, production authorization, achieved AGI/ASI or general RSI.')
page()

heading('Evidence taxonomy for every partner session','24 · Never blur authority levels')
two_col([('First-party real','REAL-001 and its matching first-party replay receipt. Supports a scoped transfer-mechanism claim.'),('Deterministic reference','COMP-001, Proof Run fixtures, Make gates and local implementation. Supports mechanism and reviewability claims.'),('Illustrative partner','Strategic Partnership Diligence and the sample Governed Decision State. Teaches the product and commercial path; not real partner evidence.'),('Independent external','Separate operator or institution reproduces the result and signs an attestation.'),('Delayed outcome','The result survives time, transfer and real operational consequences.')],('Evidence lane','Permitted use'))
quote('Theatre creates attention. Evidence determines authority.')
page()

heading('AI Council: six lenses for one objective','25 · Dynamic review')
cards([('Executive','What decision becomes faster, safer or more defensible?',GOLD),('Technical','Can the state transition be reproduced from pinned artifacts?',BLUE),('Risk / governance','Which unsupported assumption could create irreversible harm?',RED),('Commercial','Can repeated proof demand become a reusable capability asset?',GREEN),('Validator','Are reviewers genuinely independent across control and incentives?',CYAN),('RSI','Did admitted capability create persistent advantage on fresh Mission 2?',PURPLE)],3)
para('Council rule: each lens must state the proof question, required artifacts, authority consequence and the next cheapest decisive test. Models may advise; the agreed gate controls.')
page()

heading('Where partners can use GoalOS now','26 · Immediate business value')
two_col([('Strategic partnership diligence','Technical fit, commercial logic, counterparty evidence and reversible closing conditions.'),('Enterprise AI procurement','Vendor claims under a common security, data, cost and replay rubric.'),('AI product launch readiness','Connect product claims, evals, policy, incident response, rollback and release authority.'),('M&A and technology diligence','Technical, security, data, IP, operations and post-close capability map.'),('Defensive cybersecurity','Repo-owned scope, evidence-bearing findings, human review and safe remediation.'),('R&D portfolio governance','Admit only scoped, fresh, replayable results into the capability graph.')],('Mission class','Partner outcome'))
page()

heading('The commercial expansion ladder','27 · Land with proof; expand with capability')
two_col([('1. Proof Mission Sprint','One bounded objective becomes a reviewer-ready Governed Decision State.'),('2. Enterprise Proof OS','Recurring Mission Contracts, Dockets, Chronicle records and audit workflows.'),('3. Private Validated Skill Graph','Customer-specific governed capability memory with scope, replay and revocation.'),('4. Managed AGI Node Council','Independent validation without requiring the customer to operate all infrastructure.'),('5. Cross-Institution Proof Network','Jobs, validators, challenges and capability passports across organizations.')],('Offer','Customer outcome'))
quote('Sell verified outcomes first. Let repeated demand become Proof OS. Let validated capability become the durable asset.')
page()

heading('The 15-business-day Founding Proof Mission','28 · First engagement')
two_col([('Days 1–2','Freeze the decision, scope, proof level, validators, blocked claims and rollback.'),('Days 3–8','Create custom AGI Jobs, source map, ProofBundles, contradiction register, cost and risk ledgers.'),('Days 9–12','Independent review, replay, repair loop, challenge handling and public/private packaging.'),('Days 13–15','Governed Decision State, Action Graph, Chronicle candidate and fresh Mission 2 protocol.')],('Window','Deliverable'))
quote('One objective. One sponsor. One effective independent reviewer. One bounded decision window. One fresh Mission 2.')
page()

heading('Enterprise participation: serious and proportionate','29 · Receipt-bound operation')
two_col([('Authority','The accepting representative confirms professional capacity and authority to bind the organization.'),('Exclusive responsibility','The Participant remains responsible for objectives, data, people, systems, wallets, approvals, taxes, compliance and downstream use.'),('Risk and non-recourse','The Participant independently evaluates technical, AI, network and digital-asset risks.'),('Hold harmless and non-joinder','Hold harmless for Participant-caused matters and keep user-to-user matters between relevant Participants.'),('French-first evidence','French terms are made available first where required; language election is separately recorded.'),('Receipt','Organization, representative, role, language, version and exact hashes travel with the mission.'),('Montréal process','Permitted professional disputes follow the accepted notice, preservation, conference and Montréal process, subject to mandatory law.')],('Lane','Partner-facing implementation'))
para('Use the participation layer when a partner chooses to operate protected commercial functions. Do not make it the opening act.')
page()

heading('Diligence gates before production claims','30 · Partner confidence')
two_col([('Category fit','Is the proof gap material to a real decision class?'),('Mission fit','Is there one consequential but reversible objective?'),('Evidence access','Can primary evidence and controlled private appendices be made available?'),('Reviewer independence','Can a qualified reviewer be selected without correlated control?'),('Operational boundary','Are data, tool, wallet, action and authority limits explicit?'),('Security and replay','Are manifests, tests, monitoring and rollback independently reviewable?'),('Commercial path','Does repeated demand justify Proof OS, VSG or managed validation?'),('Mission 2','What fresh task would falsify the compounding hypothesis?')],('Gate','Question'))
quote('The best diligence system does not hide uncertainty. It converts uncertainty into a finite proof plan.')
page()

heading('High-caliber objections—and disciplined answers','31 · Boardroom response guide')
two_col([('“Is this just another agent platform?”','No. The core product is the proof-and-authority layer: Mission Contract, Docket, validation, Chronicle and rollback.'),('“Where is the real evidence?”','Show REAL-001 precisely, then state the independent replay and human-review debt.'),('“Why not use our existing governance?”','GoalOS can wrap existing governance with typed proof objects, replay, scoped memory and Mission 2 transfer tests.'),('“Why blockchain?”','Only for shared commitments, attestations, lineage, challenge and settlement where those properties add value; private intelligence remains controlled.'),('“Why a token?”','The partner proposition does not require token exposure. $AGIALPHA is the professional bonded-authority rail behind relevant transitions.'),('“Is this production-ready?”','The interactive institution is implemented; deployment-specific security, independent replay and production authorization remain explicit gates.'),('“What do you need from us?”','One objective, one sponsor, one reviewer, one evidence window and one decision deadline.')],('Objection','Answer'))
page()

heading('The final partnership close','32 · From masterclass to mission')
quote('Give GoalOS one consequential but reversible objective. Let the evidence determine what authority the result earns.')
bullets(['One consequential objective the partner already needs to decide.','One executive sponsor with authority to define success.','One effective independent reviewer.','One bounded, measurable and reversible decision window.','One public-safe Evidence Docket and controlled private appendix.','One Mission 2 transfer test before any compounding claim or expansion.'])
image('deal_architect.png',6.45,'The Deal Architect turns interest into a specific proof-bearing engagement.')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('AI creates output. GoalOS creates proof.\nChronicle decides what earns memory. Independent evidence determines reality.');r.font.name='Georgia';r.font.size=Pt(14);r.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY)
page()

heading('Canonical source and artifact map','Appendix · What to open when')
two_col([('Masterpiece launcher','START_HERE.html — dynamic session composer, RSI gate, AI Council and Deal Architect.'),('Masterpiece Executive Deck','32-slide executive narrative with speaker notes.'),('APEX Boardroom','Premium short-form executive theatre.'),('Grand Institution','Complete operating labs and deep curriculum.'),('Sovereign Evidence','REAL-001, replay, RSI and Mission 2 transfer.'),('Production Showcase v2.1','Live mission, Docket, Governed Decision State and participation receipt.'),('Fieldbook / facilitator corpus','Workshops, scripts, casebook, diligence dossier and charter templates.'),('Evidence directory','REAL-001 proof pack and deterministic reference fixtures.'),('Legal directory','French-first bilingual v2.1 participation instrument and manifest.'),('Source corpus','Mission OS, AEP-001, proof loops, VSG, Merkle, RSI and AGI ALPHA.')],('Artifact','Use'))
para('Release identifiers and SHA-256 checksums are supplied at the package root. The source editions remain intact for provenance and independent use.')

# Metadata
props=doc.core_properties;props.title='GoalOS Partner MASTERCLASS — Masterpiece Edition Operating Guide';props.subject='Partner presentation, operating, evidence, RSI and engagement guide';props.author='GoalOS / MONTREAL.AI';props.keywords='GoalOS, partner, masterclass, proof-carrying intelligence, RSI, Evidence Docket, Governed Decision State';props.comments='Release GOS-PM-MASTERPIECE-2026.07-v3.0'
doc.save(DOC)
print(DOC)
